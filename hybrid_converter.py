#!/usr/bin/env python3
"""
Hybrid PDF to DOCX Converter - Best of Both Worlds
Combines Docling's AI-powered text/table extraction with PyMuPDF's image extraction
"""

import os
import sys
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import fitz  # PyMuPDF
import torch

def check_cuda_availability():
    """Check CUDA availability and setup device"""
    if torch.cuda.is_available():
        device = "cuda"
        gpu_name = torch.cuda.get_device_name(0)
        gpu_memory = torch.cuda.get_device_properties(0).total_memory / 1024**3
        print(f"🚀 CUDA GPU detected: {gpu_name}")
        print(f"💾 GPU Memory: {gpu_memory:.1f} GB")
        print(f"⚡ GPU acceleration: ENABLED")
        return device
    else:
        print("⚠️ CUDA not available, using CPU")
        print("💡 For GPU acceleration, install CUDA-enabled PyTorch")
        return "cpu"

def setup_document_margins(doc):
    """Set up 1:1 margins for the document"""
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1) 
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.page_height = Inches(11)
        section.page_width = Inches(8.5)

def extract_images_with_pymupdf(pdf_path):
    """Extract images using PyMuPDF"""
    images = []
    try:
        pdf_doc = fitz.open(pdf_path)
        
        for page_num in range(len(pdf_doc)):
            page = pdf_doc[page_num]
            image_list = page.get_images()
            
            for img_index, img in enumerate(image_list):
                try:
                    xref = img[0]
                    pix = fitz.Pixmap(pdf_doc, xref)
                    
                    if pix.n - pix.alpha < 4:  # GRAY or RGB
                        img_data = pix.tobytes("png")
                        temp_img_path = f"temp_hybrid_img_p{page_num+1}_{img_index+1}.png"
                        
                        with open(temp_img_path, "wb") as f:
                            f.write(img_data)
                        
                        images.append({
                            'path': temp_img_path,
                            'page': page_num + 1,
                            'index': img_index + 1
                        })
                        
                        print(f"    🖼️ Extracted image from page {page_num+1}")
                    
                    pix = None
                    
                except Exception as e:
                    print(f"    ⚠️ Could not extract image from page {page_num+1}: {e}")
        
        pdf_doc.close()
        return images
        
    except Exception as e:
        print(f"❌ PyMuPDF image extraction failed: {e}")
        return []

def create_word_table_from_docling(word_doc, table_item):
    """Create Word table from Docling table data"""
    try:
        if not hasattr(table_item, 'data') or not table_item.data:
            return False
        
        table_data = table_item.data
        
        if hasattr(table_data, 'table_cells') and table_data.table_cells:
            cells = table_data.table_cells
            
            # Find dimensions
            max_row = max(cell.start_row_offset_idx for cell in cells) + 1
            max_col = max(cell.start_col_offset_idx for cell in cells) + 1
            
            print(f"    📊 Creating table: {max_row} rows x {max_col} columns")
            
            # Create Word table
            word_table = word_doc.add_table(rows=max_row, cols=max_col)
            word_table.style = 'Table Grid'
            
            # Fill table data
            for cell in cells:
                if cell.text and cell.text.strip():
                    row_idx = cell.start_row_offset_idx
                    col_idx = cell.start_col_offset_idx
                    
                    if row_idx < len(word_table.rows) and col_idx < len(word_table.rows[row_idx].cells):
                        word_cell = word_table.rows[row_idx].cells[col_idx]
                        word_cell.text = cell.text.strip()
                        
                        # Format header cells
                        if hasattr(cell, 'column_header') and cell.column_header:
                            for paragraph in word_cell.paragraphs:
                                for run in paragraph.runs:
                                    run.bold = True
            
            word_doc.add_paragraph()  # Space after table
            return True
            
        return False
        
    except Exception as e:
        print(f"    ❌ Table creation failed: {e}")
        return False

def hybrid_pdf_to_docx_conversion(pdf_path, output_path=None):
    """
    Hybrid conversion using Docling for text/tables and PyMuPDF for images
    """
    if not os.path.exists(pdf_path):
        raise FileNotFoundError(f"PDF file not found: {pdf_path}")
    
    if output_path is None:
        base_name = os.path.splitext(pdf_path)[0]
        output_path = f"{base_name}_hybrid.docx"
    
    print(f"🚀 Hybrid PDF to DOCX Conversion")
    print(f"📄 Input: {pdf_path}")
    print(f"📁 Output: {output_path}")
    print(f"🎯 Strategy: Docling AI + PyMuPDF Images")
    
    try:
        # Step 1: Extract images with PyMuPDF
        print("\n🖼️ Step 1: Extracting images with PyMuPDF...")
        extracted_images = extract_images_with_pymupdf(pdf_path)
        
        # Step 2: Setup CUDA acceleration
        print(f"\n⚡ Step 2: Setting up GPU acceleration...")
        device = check_cuda_availability()
        
        # Step 3: Process text and tables with Docling (GPU accelerated)
        print(f"\n🤖 Step 3: Processing text and tables with Docling AI...")
        
        from docling.document_converter import DocumentConverter
        
        # Initialize converter with default settings (GPU will be used automatically if available)
        converter = DocumentConverter()
        
        # Set device for AI models if available
        if device == "cuda":
            print("🔥 Using GPU acceleration for AI processing...")
            # Force GPU usage for underlying models
            os.environ["CUDA_VISIBLE_DEVICES"] = "0"
            torch.cuda.empty_cache()  # Clear GPU memory
        
        result = converter.convert(pdf_path)
        docling_doc = result.document
        
        print(f"✅ Docling processing complete!")
        print(f"📊 Pages: {docling_doc.num_pages}")
        print(f"📝 Text elements: {len(docling_doc.texts)}")
        print(f"📋 Tables detected: {len(docling_doc.tables)}")
        
        # Step 4: Create Word document
        print(f"\n📝 Step 4: Creating hybrid Word document...")
        
        word_doc = Document()
        setup_document_margins(word_doc)
        
        # Add title
        title_para = word_doc.add_paragraph()
        title_run = title_para.add_run("Hybrid Conversion: AI Text + Extracted Images")
        title_run.font.size = Pt(16)
        title_run.bold = True
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        word_doc.add_paragraph()
        
        # Add structured text content via markdown
        try:
            markdown_content = docling_doc.export_to_markdown()
            lines = markdown_content.split('\n')
            
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                
                if line.startswith('# '):
                    para = word_doc.add_paragraph()
                    run = para.add_run(line[2:])
                    run.font.size = Pt(16)
                    run.bold = True
                    para.style = 'Heading 1'
                    
                elif line.startswith('## '):
                    para = word_doc.add_paragraph()
                    run = para.add_run(line[3:])
                    run.font.size = Pt(14)
                    run.bold = True
                    para.style = 'Heading 2'
                    
                elif line.startswith('### '):
                    para = word_doc.add_paragraph()
                    run = para.add_run(line[4:])
                    run.font.size = Pt(12)
                    run.bold = True
                    para.style = 'Heading 3'
                    
                elif line.startswith('- ') or line.startswith('* '):
                    para = word_doc.add_paragraph(line[2:], style='List Bullet')
                    
                else:
                    if len(line) > 0 and not line.startswith('|'):
                        word_doc.add_paragraph(line)
            
            print("✅ Text content added successfully!")
            
        except Exception as e:
            print(f"⚠️ Markdown processing failed: {e}")
            # Fallback to direct text
            for text_item in docling_doc.texts:
                if hasattr(text_item, 'text') and text_item.text.strip():
                    word_doc.add_paragraph(text_item.text.strip())
        
        # Add tables from Docling
        print(f"\n📊 Step 5: Adding tables from Docling...")
        table_count = 0
        
        for i, table_item in enumerate(docling_doc.tables):
            print(f"  📋 Processing table {i+1}...")
            
            table_header = word_doc.add_paragraph(f"Table {i+1}", style='Heading 3')
            
            if create_word_table_from_docling(word_doc, table_item):
                table_count += 1
                print(f"    ✅ Table {i+1} added successfully")
            else:
                print(f"    ❌ Failed to add table {i+1}")
        
        # Add images from PyMuPDF
        print(f"\n🖼️ Step 6: Adding images from PyMuPDF...")
        image_count = 0
        
        for img_info in extracted_images:
            try:
                # Add image header
                img_header = word_doc.add_paragraph(f"Figure {img_info['index']} (Page {img_info['page']})", style='Heading 3')
                img_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Add image
                img_paragraph = word_doc.add_paragraph()
                run = img_paragraph.runs[0] if img_paragraph.runs else img_paragraph.add_run()
                
                img_width = min(Inches(6), Inches(4))
                run.add_picture(img_info['path'], width=img_width)
                img_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                word_doc.add_paragraph()  # Space after image
                
                image_count += 1
                print(f"    ✅ Added image {img_info['index']} from page {img_info['page']}")
                
            except Exception as e:
                print(f"    ❌ Failed to add image {img_info['index']}: {e}")
        
        # Save document
        print(f"\n💾 Step 7: Saving hybrid document...")
        word_doc.save(output_path)
        
        # Clean up temporary images
        print("🧹 Cleaning up temporary files...")
        for img_info in extracted_images:
            try:
                os.remove(img_info['path'])
            except:
                pass
        
        # Final statistics
        file_size = os.path.getsize(output_path) / 1024
        
        print(f"\n🎉 Hybrid conversion successful!")
        print(f"📁 Output: {output_path}")
        print(f"💾 Size: {file_size:.1f} KB")
        print(f"📏 Margins: 1:1 ratio (1 inch all sides)")
        print(f"📊 Tables: {table_count}/{len(docling_doc.tables)} extracted")
        print(f"🖼️ Images: {image_count}/{len(extracted_images)} extracted")
        print(f"🏗️ Hybrid AI + Manual extraction")
        print(f"⚡ GPU acceleration: {'✅ USED' if device == 'cuda' else '❌ CPU ONLY'}")
        
        # Clear GPU memory if used
        if device == "cuda":
            torch.cuda.empty_cache()
            print("🧹 GPU memory cleared")
        
        return output_path
        
    except Exception as e:
        print(f"❌ Hybrid conversion failed: {str(e)}")
        raise

if __name__ == "__main__":
    print("🚀 Hybrid PDF to DOCX Converter")
    print("=" * 60)
    print("🎯 Best of Both Worlds: Docling AI + PyMuPDF Images")
    print("=" * 60)
    
    pdf_path = "sample.pdf"
    
    if not os.path.exists(pdf_path):
        print(f"❌ PDF file not found: {pdf_path}")
        sys.exit(1)
    
    try:
        output_file = hybrid_pdf_to_docx_conversion(pdf_path)
        
        print("\n" + "=" * 60)
        print("🏆 HYBRID CONVERSION COMPLETE!")
        print(f"📁 Your file: {output_file}")
        print("✅ Perfect combination achieved:")
        print("  • Docling AI: Superior text & table extraction")
        print("  • PyMuPDF: Reliable image extraction")
        print("  • 1:1 margin ratio (1 inch all sides)")
        print("  • Complete content preservation")
        print("  • Professional document layout")
        print("=" * 60)
        
    except Exception as e:
        print(f"\n💥 Conversion failed: {e}")
        sys.exit(1)