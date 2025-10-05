import os
from typing import List, Tuple
import tempfile
import io

import gradio as gr
from PIL import Image
from PyPDF2 import PdfReader

from langchain_chroma import Chroma
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_openai import ChatOpenAI
from langchain.chains import RetrievalQA
from langchain.docstore.document import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter

# Optional deps
try:
	import docx  # python-docx
except Exception:
	docx = None
try:
	import whisper  # openai-whisper
except Exception:
	whisper = None
try:
	import pytesseract
except Exception:
	pytesseract = None

OPENAI_API_KEY = os.getenv("OPENAI_API_KEY", "")
OPENAI_API_BASE = os.getenv("OPENAI_API_BASE", "https://openrouter.ai/api/v1")

PERSIST_DIR = os.getenv("PERSIST_DIR", "./db")
VISION_MODEL = os.getenv("VISION_MODEL", "qwen/qwen-2.5-vl-72b-instruct")

# ---------------------------
# Helpers for extraction
# ---------------------------

def extract_text_from_pdf_bytes(data: bytes) -> str:
	text = ""
	try:
		with tempfile.NamedTemporaryFile(suffix=".pdf", delete=True) as tf:
			tf.write(data)
			tf.flush()
			reader = PdfReader(tf.name)
			for page in reader.pages:
				text += page.extract_text() or ""
	except Exception:
		pass
	return text


def extract_text_from_docx_bytes(data: bytes) -> str:
	if docx is None:
		return ""
	try:
		with tempfile.NamedTemporaryFile(suffix=".docx", delete=True) as tf:
			tf.write(data)
			tf.flush()
			d = docx.Document(tf.name)
			return "\n".join(p.text for p in d.paragraphs)
	except Exception:
		return ""


def transcribe_audio_bytes(data: bytes, original_name: str = "", whisper_size: str = "base") -> str:
	if whisper is None:
		return ""
	try:
		ext = os.path.splitext(original_name)[1].lower() or ".wav"
		with tempfile.NamedTemporaryFile(suffix=ext, delete=True) as tf:
			tf.write(data)
			tf.flush()
			model = whisper.load_model(whisper_size)
			result = model.transcribe(tf.name)
			return result.get("text", "")
	except Exception:
		return ""


def ocr_image(img: Image.Image) -> str:
	if pytesseract is None:
		return ""
	try:
		return pytesseract.image_to_string(img)
	except Exception:
		return ""


# ---------------------------
# Build RAG from uploaded files
# ---------------------------

def _read_file_input(f) -> Tuple[str, bytes]:
	"""Return (lowercased name/path, bytes) from Gradio File input.
	Prefer filesystem path from string/NamedString or attributes; fall back to .read() if needed.
	"""
	# String or path-like
	if isinstance(f, (str, os.PathLike)):
		path = str(f)
		try:
			with open(path, "rb") as fh:
				return path.lower(), fh.read()
		except Exception:
			return path.lower(), b""
	# Objects (e.g., NamedString or UploadedFile)
	path = getattr(f, "name", None) or getattr(f, "path", None)
	if isinstance(path, (str, os.PathLike)):
		try:
			with open(str(path), "rb") as fh:
				return str(path).lower(), fh.read()
		except Exception:
			pass
	# Last resort: file-like read
	if hasattr(f, "read"):
		try:
			data = f.read()
			return (str(getattr(f, "name", "uploaded"))).lower(), (data or b"")
		except Exception:
			return (str(getattr(f, "name", "uploaded"))).lower(), b""
	return "", b""


def build_rag_from_files(files: List, whisper_size: str = "base", api_key: str = "") -> Tuple[RetrievalQA, List[Image.Image]]:
	texts: List[str] = []
	cached_images: List[Image.Image] = []

	for f in files or []:
		name, data = _read_file_input(f)
		if not data:
			continue
		if name.endswith(".pdf"):
			texts.append(extract_text_from_pdf_bytes(data))
		elif name.endswith(".txt"):
			try:
				texts.append(data.decode("utf-8", errors="ignore"))
			except Exception:
				pass
		elif name.endswith(".docx"):
			texts.append(extract_text_from_docx_bytes(data))
		elif name.endswith((".png", ".jpg", ".jpeg")):
			try:
				img = Image.open(io.BytesIO(data))
				cached_images.append(img)
				ocr_txt = ocr_image(img)
				if ocr_txt.strip():
					texts.append(ocr_txt)
			except Exception:
				pass
		elif name.endswith((".wav", ".mp3", ".m4a")):
			tr = transcribe_audio_bytes(data, original_name=name, whisper_size=whisper_size)
			if tr.strip():
				texts.append(tr)

	documents = [Document(page_content=t) for t in texts if t and t.strip()]

	embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")
	if documents:
		splitter = RecursiveCharacterTextSplitter(chunk_size=800, chunk_overlap=120, separators=["\n\n", "\n", ". ", ".", " "])
		documents = splitter.split_documents(documents)
		vectordb = Chroma.from_documents(documents, embedding=embeddings, persist_directory=PERSIST_DIR)
	else:
		vectordb = Chroma(embedding_function=embeddings, persist_directory=PERSIST_DIR)

	key = (api_key or OPENAI_API_KEY or "").strip()
	if not key:
		raise ValueError("OpenRouter API key is not set. Enter it in the UI.")

	text_llm = ChatOpenAI(
		model="openai/gpt-oss-120b",
		temperature=0,
		max_tokens=512,
		openai_api_key=key,
		openai_api_base=OPENAI_API_BASE,
	)

	retriever = vectordb.as_retriever(search_type="mmr", search_kwargs={"k": 4, "fetch_k": 24})
	qa = RetrievalQA.from_chain_type(llm=text_llm, retriever=retriever, chain_type="stuff")
	return qa, cached_images


# ---------------------------
# Vision helper
# ---------------------------

def analyze_images_with_vision(images: List[Image.Image], prompt: str, api_key: str = "") -> str:
	if not images:
		return ""
	try:
		key = (api_key or OPENAI_API_KEY or "").strip()
		if not key:
			return ""
		vision_llm = ChatOpenAI(
			model=VISION_MODEL,
			temperature=0,
			max_tokens=512,
			openai_api_key=key,
			openai_api_base=OPENAI_API_BASE,
		)
		parts = []
		for idx, img in enumerate(images):
			message = [{"role": "user", "content": prompt, "image": img}]
			resp = vision_llm.invoke(message)
			parts.append(f"Image {idx+1}: " + (getattr(resp, "content", "") or str(resp)))
		return "\n\n".join(parts)
	except Exception as e:
		return f"[vision error] {e}"


# ---------------------------
# Gradio UI
# ---------------------------

def on_upload(files, whisper_size, api_key_state):
	qa, images = build_rag_from_files(files, whisper_size=whisper_size, api_key=api_key_state or "")
	return qa, images, gr.update(visible=True), "✅ Files ingested. Ask your question."


def on_chat(message, history, qa_state, images, analyze_images, api_key_state):
	if qa_state is None:
		new_history = list(history or []) + [
			{"role": "user", "content": message},
			{"role": "assistant", "content": "Please upload files first and set API key."},
		]
		return new_history, qa_state
	answer = ""
	try:
		result = qa_state.invoke({"query": message})
		answer = result.get("result", "")
	except Exception as e:
		answer = f"Error: {e}"

	# Optional vision analysis when requested
	do_vision = analyze_images or (message.lower().startswith("image:"))
	vision_text = ""
	if do_vision and images:
		q = message.split(":", 1)[1].strip() if message.lower().startswith("image:") else message
		vision_text = analyze_images_with_vision(images, q, api_key=api_key_state or "")
		if vision_text:
			answer = (answer + "\n\n" + vision_text).strip()

	new_history = list(history or []) + [
		{"role": "user", "content": message},
		{"role": "assistant", "content": answer},
	]
	return new_history, qa_state


def build_app():
	with gr.Blocks(title="RAG Chat") as demo:
		gr.Markdown("## Chat with your files (PDF, TXT, DOCX, Images, Audio)")
		state_qa = gr.State(value=None)  # stores RetrievalQA
		state_images = gr.State(value=[])  # stores cached images
		state_api_key = gr.State(value=OPENAI_API_KEY)

		with gr.Row():
			api_key_box = gr.Textbox(label="OpenRouter API Key", value=OPENAI_API_KEY, type="password", placeholder="sk-or-...", scale=3)
			apply_key = gr.Button("Set Key", scale=1)
			key_status = gr.Markdown(visible=False)

		with gr.Row():
			files = gr.File(label="Upload files", file_count="multiple", type="filepath")
			whisper_size = gr.Dropdown(choices=["tiny","base","small","medium","large"], value="base", label="Whisper model")
			analyze_images = gr.Checkbox(label="Analyze images automatically", value=False)
			upload_btn = gr.Button("Ingest")

		status = gr.Markdown(visible=False)

		chatbot = gr.Chatbot(type="messages", height=420)
		msg = gr.Textbox(label="Message", placeholder="Ask a question about your files...")
		send = gr.Button("Send")

		def set_key(k):
			masked = (k[:6] + "..." + k[-4:]) if k and len(k) > 10 else (k or "<empty>")
			return k, gr.update(visible=True, value=f"🔑 Using API key: {masked}")

		apply_key.click(set_key, inputs=[api_key_box], outputs=[state_api_key, key_status])

		upload_btn.click(on_upload, inputs=[files, whisper_size, state_api_key], outputs=[state_qa, state_images, status, status])
		send.click(on_chat, inputs=[msg, chatbot, state_qa, state_images, analyze_images, state_api_key], outputs=[chatbot, state_qa])
		msg.submit(on_chat, inputs=[msg, chatbot, state_qa, state_images, analyze_images, state_api_key], outputs=[chatbot, state_qa])
	return demo


if __name__ == "__main__":
	app = build_app()
	app.launch(server_name="0.0.0.0", server_port=7860)
